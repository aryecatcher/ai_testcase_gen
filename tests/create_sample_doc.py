from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_sample_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('在线书店系统 - 购物车与结算模块需求说明书', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 1
    doc.add_heading('1. 概述', level=1)
    doc.add_paragraph('本模块主要负责用户的商品加购、购物车管理、优惠券应用及最终订单结算功能。')

    # Section 2
    doc.add_heading('2. 详细功能需求', level=1)

    # 2.1 Add to Cart
    doc.add_heading('2.1 添加购物车', level=2)
    p = doc.add_paragraph()
    p.add_run('功能描述：').bold = True
    p.add_run(' 用户在商品详情页点击“加入购物车”。')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '规则项'
    hdr_cells[1].text = '规则描述'
    hdr_cells[2].text = '预期反馈'
    
    rules = [
        ('库存检查', '若商品库存不足', '提示“库存不足”，禁止加入'),
        ('限购检查', '若商品属于限购商品（每人限购2件）且购物车内已存在2件', '提示“已达限购数量”'),
        ('数量限制', '单次添加数量必须 > 0 且 <= 99', '输入非法数字时自动重置为1或99'),
        ('重复添加', '若购物车已存在该商品', '数量累加，不新增记录')
    ]
    
    for item, desc, result in rules:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc
        row_cells[2].text = result

    # 2.2 Coupon
    doc.add_heading('2.2 优惠券应用', level=2)
    doc.add_paragraph('用户在结算页选择优惠券，系统自动计算抵扣金额。')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('满减券：').bold = True
    ul.add_run(' 订单金额满 100 减 10。若应用后订单金额 < 0，则实付金额为 0（虽然逻辑上不太可能，但需校验）。')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('有效期校验：').bold = True
    ul.add_run(' 只能选择当前时间在 [生效时间, 过期时间] 范围内的优惠券。')

    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('互斥规则：').bold = True
    ul.add_run(' “满减券”与“折扣券”不可叠加使用。')

    # 2.3 Checkout
    doc.add_heading('2.3 提交订单', level=2)
    doc.add_paragraph('提交订单时需进行最终校验：')
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '校验点'
    hdr_cells2[1].text = '逻辑'
    
    checks = [
        ('收货地址', '地址ID必须有效且属于当前用户'),
        ('价格变动', '若结算时商品价格发生变动（与加购时不一致），提示用户并刷新价格'),
        ('库存扣减', '提交成功后，预扣减数据库库存；若扣减失败则回滚事务')
    ]
    
    for point, logic in checks:
        row_cells = table2.add_row().cells
        row_cells[0].text = point
        row_cells[1].text = logic

    # Save
    file_path = 'sample_requirements.docx'
    doc.save(file_path)
    print(f"File created at: {file_path}")

if __name__ == "__main__":
    create_sample_doc()
