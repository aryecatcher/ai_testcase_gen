import docx
import pandas as pd
import json
from typing import List, Dict, Any
from pathlib import Path
from loguru import logger

class DocParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_type = Path(file_path).suffix.lower()

    def parse(self) -> str:
        """
        解析文档，返回合并后的全文本。
        后续交给 AI 做结构化提取。
        """
        logger.info(f"Parsing file: {self.file_path}")
        try:
            if self.file_type == ".docx":
                return self._parse_docx()
            elif self.file_type == ".xlsx":
                return self._parse_excel()
            else:
                raise ValueError(f"Unsupported file type: {self.file_type}")
        except Exception as e:
            logger.error(f"Error parsing file: {e}")
            raise

    def _parse_docx(self) -> str:
        doc = docx.Document(self.file_path)
        full_text = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # 提取表格
        for table in doc.tables:
            full_text.append("\n[Table Start]")
            # 获取表头（假设第一行是表头）
            headers = []
            if len(table.rows) > 0:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                full_text.append(" | ".join(headers))
                full_text.append("-" * len(" | ".join(headers)))
            
            # 获取数据行
            for row in table.rows[1:]:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(" | ".join(row_text))
            full_text.append("[Table End]\n")
                
        return "\n".join(full_text)

    def _parse_excel(self) -> str:
        # Excel 读取所有 Sheet，转为 CSV 格式的字符串
        dfs = pd.read_excel(self.file_path, sheet_name=None)
        text_parts = []
        for sheet_name, df in dfs.items():
            text_parts.append(f"--- Sheet: {sheet_name} ---")
            # 将 DataFrame 转换为 CSV 字符串
            text_parts.append(df.to_csv(index=False))
        return "\n\n".join(text_parts)
